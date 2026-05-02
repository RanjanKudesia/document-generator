import base64
import logging
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from lxml import etree

from app.schemas.document_generation_schema import (
    DocumentGenerationRequest,
    ExtractedData,
    ExtractedDocumentDefaults,
    ExtractedParagraph,
    ExtractedTable,
    ExtractedXmlParagraph,
    ExtractedXmlRun,
    ExtractedXmlData,
    ExtractedMediaItem,
    ExtractedStyle,
    ParagraphBlock,
    TableBlock,
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_NS = {"w": W_NS}
W_VAL_LITERAL = "w:val"


class DocxGenerationPipeline:
    def __init__(self) -> None:
        self.logger = logging.getLogger(__name__)

    def run(self, payload: DocumentGenerationRequest, file_name: str) -> bytes:
        _ = file_name
        document = Document()

        if payload.extracted_data is not None:
            if isinstance(payload.extracted_data, ExtractedXmlData):
                self._add_xml_extracted_payload(
                    document, payload.extracted_data)
            else:
                self._add_extracted_payload(document, payload.extracted_data)
        else:
            if payload.title:
                document.add_heading(payload.title, level=1)

            for block in payload.blocks:
                if isinstance(block, ParagraphBlock):
                    self._add_paragraph_block(document, block)
                elif isinstance(block, TableBlock):
                    self._add_table_block(document, block)

        output = BytesIO()
        document.save(output)
        return output.getvalue()

    def _add_xml_extracted_payload(self, document: Document, extracted_data: ExtractedXmlData) -> None:
        self._apply_document_defaults(
            document, extracted_data.document_defaults)
        self._apply_extracted_styles(document, extracted_data.styles)

        if extracted_data.parsed_body:
            self._add_xml_parsed_body(document, extracted_data)
            return

        doc_part = next(
            (p for p in extracted_data.parts if p.path == "word/document.xml"), None)
        if doc_part is None or not doc_part.xml.strip():
            return

        try:
            root = etree.fromstring(doc_part.xml.encode("utf-8"))
        except Exception:
            return

        body = root.find("w:body", XML_NS)
        if body is None:
            return

        for child in body:
            local = child.tag.rsplit(
                "}", 1)[-1] if "}" in child.tag else child.tag
            if local == "p":
                self._add_xml_paragraph(document, child)
            elif local == "tbl":
                self._add_xml_table(document, child)

    def _add_xml_parsed_body(self, document: Document, extracted_data: ExtractedXmlData) -> None:
        for item in extracted_data.parsed_body:
            if item.type == "paragraph" and item.paragraph is not None:
                self._add_xml_parsed_paragraph(
                    document, item.paragraph, extracted_data.relationships)
            elif item.type == "table" and item.table is not None:
                self._add_xml_parsed_table(
                    document, item.table, extracted_data.relationships)

    def _add_xml_parsed_paragraph(self, document: Document, paragraph_data: ExtractedXmlParagraph, relationships: dict[str, str]) -> None:
        style_name = self._resolve_xml_paragraph_style_name(paragraph_data)
        if style_name:
            try:
                paragraph = document.add_paragraph(style=style_name)
            except KeyError:
                paragraph = document.add_paragraph()
        else:
            paragraph = document.add_paragraph()

        alignment = self._map_xml_alignment(paragraph_data.alignment)
        if alignment is not None:
            paragraph.alignment = alignment

        if paragraph_data.runs:
            for run in paragraph_data.runs:
                target = run.hyperlink_target
                if not target and run.hyperlink_rid:
                    target = relationships.get(run.hyperlink_rid)

                if target or run.hyperlink_anchor:
                    self._add_xml_hyperlink_run(
                        paragraph,
                        run,
                        target,
                        run.hyperlink_anchor,
                    )
                else:
                    docx_run = paragraph.add_run(run.text or "")
                    self._apply_xml_run_formatting(docx_run, run)
                    for media_item in run.embedded_media:
                        self._add_media_to_paragraph(paragraph, media_item)
        elif paragraph_data.text:
            paragraph.add_run(paragraph_data.text)

    def _add_xml_parsed_table(self, document: Document, table_data, relationships: dict[str, str]) -> None:
        rows = table_data.rows or []
        if not rows:
            return

        col_count = max((len(r.cells) for r in rows), default=0)
        if col_count == 0:
            return

        table = document.add_table(rows=len(rows), cols=col_count)
        for r_i, row in enumerate(rows):
            for c_i in range(col_count):
                cell = table.cell(r_i, c_i)
                if c_i >= len(row.cells):
                    cell.text = ""
                    continue

                src_cell = row.cells[c_i]
                # clear default empty paragraph so we can rehydrate formatting
                if cell.paragraphs:
                    first = cell.paragraphs[0]
                    first._element.getparent().remove(first._element)

                if src_cell.paragraphs:
                    for p in src_cell.paragraphs:
                        self._add_xml_parsed_paragraph_to_cell(
                            cell, p, relationships)
                else:
                    cell.text = src_cell.text or ""

    def _add_xml_parsed_paragraph_to_cell(self, cell, paragraph_data: ExtractedXmlParagraph, relationships: dict[str, str]) -> None:
        style_name = self._resolve_xml_paragraph_style_name(paragraph_data)
        if style_name:
            try:
                paragraph = cell.add_paragraph(style=style_name)
            except KeyError:
                paragraph = cell.add_paragraph()
        else:
            paragraph = cell.add_paragraph()

        alignment = self._map_xml_alignment(paragraph_data.alignment)
        if alignment is not None:
            paragraph.alignment = alignment

        if paragraph_data.runs:
            for run in paragraph_data.runs:
                target = run.hyperlink_target
                if not target and run.hyperlink_rid:
                    target = relationships.get(run.hyperlink_rid)
                if target or run.hyperlink_anchor:
                    self._add_xml_hyperlink_run(
                        paragraph,
                        run,
                        target,
                        run.hyperlink_anchor,
                    )
                else:
                    docx_run = paragraph.add_run(run.text or "")
                    self._apply_xml_run_formatting(docx_run, run)
                    for media_item in run.embedded_media:
                        self._add_media_to_paragraph(paragraph, media_item)
        elif paragraph_data.text:
            paragraph.add_run(paragraph_data.text)

    def _apply_xml_run_formatting(self, run, run_data: ExtractedXmlRun) -> None:
        if run_data.bold is not None:
            run.bold = run_data.bold
        if run_data.italic is not None:
            run.italic = run_data.italic
        if run_data.underline is not None:
            run.underline = run_data.underline
        if run_data.font_name:
            run.font.name = run_data.font_name
        if run_data.font_size_pt is not None and run_data.font_size_pt > 0:
            run.font.size = Pt(run_data.font_size_pt)
        if run_data.color_rgb:
            try:
                hex_str = run_data.color_rgb.replace("#", "").strip()
                run.font.color.rgb = RGBColor(
                    int(hex_str[0:2], 16),
                    int(hex_str[2:4], 16),
                    int(hex_str[4:6], 16),
                )
            except Exception:
                pass

    def _add_xml_hyperlink_run(
        self,
        paragraph,
        run_data: ExtractedXmlRun,
        target: str | None,
        anchor: str | None,
    ) -> None:
        hyperlink = OxmlElement("w:hyperlink")

        if anchor:
            hyperlink.set(qn("w:anchor"), anchor)
        elif target:
            try:
                r_id = paragraph.part.relate_to(
                    target, RT.HYPERLINK, is_external=True)
                hyperlink.set(qn("r:id"), r_id)
            except Exception:
                run = paragraph.add_run(run_data.text or "")
                self._apply_xml_run_formatting(run, run_data)
                return
        else:
            run = paragraph.add_run(run_data.text or "")
            self._apply_xml_run_formatting(run, run_data)
            return

        run_elem = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")

        if run_data.bold:
            rpr.append(OxmlElement("w:b"))
        if run_data.italic:
            rpr.append(OxmlElement("w:i"))
        if run_data.underline is not False:
            u_elem = OxmlElement("w:u")
            u_elem.set(qn("w:val"), "single")
            rpr.append(u_elem)

        color_hex = run_data.color_rgb.replace(
            "#", "").strip() if run_data.color_rgb else "0563C1"
        color_elem = OxmlElement("w:color")
        color_elem.set(qn("w:val"), color_hex)
        rpr.append(color_elem)

        if run_data.font_name:
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), run_data.font_name)
            fonts.set(qn("w:hAnsi"), run_data.font_name)
            rpr.append(fonts)

        if run_data.font_size_pt and run_data.font_size_pt > 0:
            half_pts = str(int(run_data.font_size_pt * 2))
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), half_pts)
            rpr.append(sz)

        run_elem.append(rpr)
        self._append_text_to_oxml_run(run_elem, run_data.text or "")
        hyperlink.append(run_elem)
        paragraph._p.append(hyperlink)

    def _append_text_to_oxml_run(self, run_elem, text: str) -> None:
        buf: list[str] = []

        def flush_text() -> None:
            if not buf:
                return
            content = "".join(buf)
            buf.clear()
            t = OxmlElement("w:t")
            t.text = content
            if content.startswith(" ") or content.endswith(" "):
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            run_elem.append(t)

        for ch in text:
            if ch == "\n":
                flush_text()
                run_elem.append(OxmlElement("w:br"))
            elif ch == "\t":
                flush_text()
                run_elem.append(OxmlElement("w:tab"))
            else:
                buf.append(ch)

        flush_text()

    def _resolve_xml_paragraph_style_name(self, paragraph_data: ExtractedXmlParagraph) -> str | None:
        if paragraph_data.is_numbered:
            return "List Number"
        if paragraph_data.is_bullet:
            return "List Bullet"
        mapped = self._map_xml_style_id(paragraph_data.style_id)
        if mapped:
            return mapped
        return None

    def _map_xml_style_id(self, style_id: str | None) -> str | None:
        if not style_id:
            return None
        mapping = {
            "Heading1": "Heading 1",
            "Heading2": "Heading 2",
            "Heading3": "Heading 3",
            "Heading4": "Heading 4",
            "Heading5": "Heading 5",
            "Heading6": "Heading 6",
            "Heading7": "Heading 7",
            "Heading8": "Heading 8",
            "Heading9": "Heading 9",
            "ListBullet": "List Bullet",
            "ListNumber": "List Number",
        }
        return mapping.get(style_id, style_id)

    def _map_xml_alignment(self, raw: str | None) -> WD_ALIGN_PARAGRAPH | None:
        if not raw:
            return None
        normalized = raw.strip().lower()
        if normalized in {"left", "start"}:
            return WD_ALIGN_PARAGRAPH.LEFT
        if normalized == "center":
            return WD_ALIGN_PARAGRAPH.CENTER
        if normalized in {"right", "end"}:
            return WD_ALIGN_PARAGRAPH.RIGHT
        if normalized in {"both", "distribute", "justify"}:
            return WD_ALIGN_PARAGRAPH.JUSTIFY
        return None

    def _add_xml_paragraph(self, document: Document, paragraph_el) -> None:
        style_name = self._xml_paragraph_style_name(paragraph_el)
        if style_name:
            try:
                paragraph = document.add_paragraph(style=style_name)
            except KeyError:
                paragraph = document.add_paragraph()
        else:
            paragraph = document.add_paragraph()

        text_chunks: list[str] = []
        for t in paragraph_el.xpath(".//w:t", namespaces=XML_NS):
            text_chunks.append(t.text or "")

        text = "".join(text_chunks)
        if text:
            paragraph.add_run(text)

    def _xml_paragraph_style_name(self, paragraph_el) -> str | None:
        pstyle = paragraph_el.find("w:pPr/w:pStyle", XML_NS)
        if pstyle is None:
            return None

        raw = pstyle.get(f"{{{W_NS}}}val")
        if not raw:
            return None

        # Map common XML style IDs to python-docx style display names.
        mapping = {
            "Heading1": "Heading 1",
            "Heading2": "Heading 2",
            "Heading3": "Heading 3",
            "Heading4": "Heading 4",
            "Heading5": "Heading 5",
            "Heading6": "Heading 6",
            "Heading7": "Heading 7",
            "Heading8": "Heading 8",
            "Heading9": "Heading 9",
            "ListBullet": "List Bullet",
            "ListNumber": "List Number",
        }
        return mapping.get(raw, raw)

    def _add_xml_table(self, document: Document, table_el) -> None:
        rows = table_el.findall("w:tr", XML_NS)
        if not rows:
            return

        col_count = 0
        for row in rows:
            cells = row.findall("w:tc", XML_NS)
            col_count = max(col_count, len(cells))

        if col_count == 0:
            return

        table = document.add_table(rows=len(rows), cols=col_count)
        for r_i, row in enumerate(rows):
            cells = row.findall("w:tc", XML_NS)
            for c_i in range(col_count):
                if c_i >= len(cells):
                    table.cell(r_i, c_i).text = ""
                    continue

                cell_text = "".join(
                    (t.text or "") for t in cells[c_i].xpath(".//w:t", namespaces=XML_NS)
                )
                table.cell(r_i, c_i).text = cell_text

    def _add_extracted_payload(self, document: Document, extracted_data: ExtractedData) -> None:
        self._apply_document_defaults(
            document, extracted_data.document_defaults)
        self._apply_extracted_styles(document, extracted_data.styles)
        self._apply_section_settings(document, extracted_data.sections)

        paragraph_by_index = {
            item.index: item for item in extracted_data.paragraphs}
        table_by_index = {item.index: item for item in extracted_data.tables}
        current_page_index: int | None = None

        if extracted_data.document_order:
            for order_item in extracted_data.document_order:
                if order_item.type == "paragraph":
                    paragraph = paragraph_by_index.get(order_item.index)
                    if paragraph is not None:
                        next_page = getattr(paragraph, "page_index", None)
                        current_page_index = self._maybe_add_page_break(
                            document, current_page_index, next_page)
                        self._add_extracted_paragraph(document, paragraph)
                elif order_item.type == "table":
                    table = table_by_index.get(order_item.index)
                    if table is not None:
                        self._add_extracted_table(document, table)
        else:
            for paragraph in sorted(extracted_data.paragraphs, key=lambda item: item.index):
                self._add_extracted_paragraph(document, paragraph)
            for table in sorted(extracted_data.tables, key=lambda item: item.index):
                self._add_extracted_table(document, table)

    def _maybe_add_page_break(
        self,
        document: Document,
        current_page: int | None,
        next_page: int | None,
    ) -> int | None:
        """Insert a page break paragraph when the page index advances."""
        if next_page is None:
            return current_page
        if current_page is None:
            return next_page
        if next_page > current_page:
            document.add_page_break()
            return next_page
        return current_page

    def _apply_section_settings(self, document: Document, sections: list) -> None:
        """Apply page dimensions and margins from the first extracted section."""
        if not sections:
            return
        try:
            from docx.shared import Twips
            first = sections[0]
            if not isinstance(first, dict):
                return
            sec = document.sections[0]
            if first.get("page_width_twips"):
                sec.page_width = Twips(first["page_width_twips"])
            if first.get("page_height_twips"):
                sec.page_height = Twips(first["page_height_twips"])
            if first.get("left_margin_twips"):
                sec.left_margin = Twips(first["left_margin_twips"])
            if first.get("right_margin_twips"):
                sec.right_margin = Twips(first["right_margin_twips"])
            if first.get("top_margin_twips"):
                sec.top_margin = Twips(first["top_margin_twips"])
            if first.get("bottom_margin_twips"):
                sec.bottom_margin = Twips(first["bottom_margin_twips"])
        except Exception:
            pass

    def _apply_document_defaults(
        self,
        document: Document,
        defaults: ExtractedDocumentDefaults | None,
    ) -> None:
        """Apply source docDefaults to base styles used by body paragraphs and inherited runs."""
        if defaults is None:
            return

        for style_name in ("Normal", "Default Paragraph Font"):
            try:
                style_obj = document.styles[style_name]
            except KeyError:
                continue

            if defaults.font_name:
                style_obj.font.name = defaults.font_name
            if defaults.font_size_pt is not None and defaults.font_size_pt > 0:
                style_obj.font.size = Pt(defaults.font_size_pt)
            if defaults.color_rgb:
                try:
                    hex_str = defaults.color_rgb.replace("#", "").strip()
                    style_obj.font.color.rgb = RGBColor(
                        int(hex_str[0:2], 16),
                        int(hex_str[2:4], 16),
                        int(hex_str[4:6], 16),
                    )
                except Exception:
                    pass

    def _apply_extracted_styles(self, document: Document, styles: list[ExtractedStyle]) -> None:
        """Apply extracted style font defaults so inherited run formatting is preserved."""
        for style_data in styles:
            if style_data.font is None:
                continue

            style_obj = self._get_or_create_style(document, style_data)
            if style_obj is None:
                continue

            font_data = style_data.font
            if font_data.name:
                style_obj.font.name = font_data.name
            else:
                self._clear_style_rpr_override(style_obj, "rFonts")
            if font_data.size_pt is not None and font_data.size_pt > 0:
                style_obj.font.size = Pt(font_data.size_pt)
            else:
                self._clear_style_rpr_override(style_obj, "sz")
                self._clear_style_rpr_override(style_obj, "szCs")
            if font_data.bold is not None:
                style_obj.font.bold = font_data.bold
            else:
                self._clear_style_rpr_override(style_obj, "b")
            if font_data.italic is not None:
                style_obj.font.italic = font_data.italic
            else:
                self._clear_style_rpr_override(style_obj, "i")
            if font_data.underline is not None:
                style_obj.font.underline = font_data.underline
            else:
                self._clear_style_rpr_override(style_obj, "u")
            if font_data.color_rgb:
                try:
                    hex_str = font_data.color_rgb.replace("#", "").strip()
                    style_obj.font.color.rgb = RGBColor(
                        int(hex_str[0:2], 16),
                        int(hex_str[2:4], 16),
                        int(hex_str[4:6], 16),
                    )
                except Exception:
                    pass
            else:
                self._clear_style_rpr_override(style_obj, "color")
            if font_data.highlight_color:
                try:
                    style_obj.font.highlight_color = WD_COLOR_INDEX[
                        font_data.highlight_color.upper()
                    ]
                except (KeyError, AttributeError):
                    pass
            else:
                self._clear_style_rpr_override(style_obj, "highlight")

    def _clear_style_rpr_override(self, style_obj, tag_name: str) -> None:
        """Remove direct run-property override from style XML so value can inherit."""
        try:
            style_el = style_obj.element
            rpr = style_el.find(qn("w:rPr"))
            if rpr is None:
                return

            child = rpr.find(qn(f"w:{tag_name}"))
            if child is not None:
                rpr.remove(child)
        except Exception:
            return

    def _get_or_create_style(self, document: Document, style_data: ExtractedStyle):
        style_name = style_data.name
        style_id = style_data.style_id

        if style_name:
            try:
                return document.styles[style_name]
            except KeyError:
                pass

        if style_id:
            try:
                return document.styles[style_id]
            except KeyError:
                pass

        if not style_name:
            return None

        style_type = (style_data.type or "").upper()
        if "PARAGRAPH" in style_type:
            create_type = WD_STYLE_TYPE.PARAGRAPH
        elif "CHARACTER" in style_type:
            create_type = WD_STYLE_TYPE.CHARACTER
        elif "TABLE" in style_type:
            create_type = WD_STYLE_TYPE.TABLE
        else:
            create_type = WD_STYLE_TYPE.PARAGRAPH

        try:
            return document.styles.add_style(style_name, create_type)
        except Exception:
            return None

    def _add_extracted_paragraph(self, document: Document, paragraph_data: ExtractedParagraph) -> None:
        paragraph = self._create_output_paragraph(document, paragraph_data)
        self._apply_paragraph_spacing(paragraph, paragraph_data)
        self._apply_paragraph_rtl(paragraph, paragraph_data)
        self._populate_output_paragraph(paragraph, paragraph_data)

    def _apply_paragraph_spacing(self, paragraph, paragraph_data: ExtractedParagraph) -> None:
        """Apply space_before_pt, space_after_pt, and line_spacing from extracted data."""
        fmt = paragraph.paragraph_format
        try:
            if paragraph_data.space_before_pt is not None and paragraph_data.space_before_pt >= 0:
                fmt.space_before = Pt(paragraph_data.space_before_pt)
            if paragraph_data.space_after_pt is not None and paragraph_data.space_after_pt >= 0:
                fmt.space_after = Pt(paragraph_data.space_after_pt)
            if paragraph_data.line_spacing is not None and paragraph_data.line_spacing > 0:
                fmt.line_spacing = paragraph_data.line_spacing
        except Exception:
            pass

    def _apply_paragraph_rtl(self, paragraph, paragraph_data: ExtractedParagraph) -> None:
        """Insert <w:bidi/> into pPr when paragraph direction is RTL."""
        if getattr(paragraph_data, "direction", None) != "rtl":
            return
        try:
            p_pr = paragraph._p.get_or_add_pPr()
            bidi = p_pr.find(qn("w:bidi"))
            if bidi is None:
                bidi = OxmlElement("w:bidi")
                p_pr.append(bidi)
        except Exception:
            pass

    def _populate_output_paragraph(self, paragraph, paragraph_data: ExtractedParagraph) -> None:
        """Apply paragraph style, alignment, and runs/text to an existing paragraph."""
        style_name = self._resolve_paragraph_style_name(paragraph_data)
        if style_name:
            try:
                paragraph.style = style_name
            except KeyError:
                pass

        alignment = self._map_alignment(paragraph_data.alignment)
        if alignment is not None:
            paragraph.alignment = alignment

        if paragraph_data.runs:
            self._add_paragraph_runs(paragraph, paragraph_data)
            return
        if paragraph_data.text:
            paragraph.add_run(paragraph_data.text)

    def _create_output_paragraph(self, document: Document, paragraph_data: ExtractedParagraph):
        """Create paragraph with best-effort style assignment."""
        style_name = self._resolve_paragraph_style_name(paragraph_data)
        if style_name:
            try:
                return document.add_paragraph(style=style_name)
            except KeyError:
                return document.add_paragraph()
        return document.add_paragraph()

    def _add_paragraph_runs(self, paragraph, paragraph_data: ExtractedParagraph) -> None:
        """Add runs and embedded media to paragraph."""
        for run_data in paragraph_data.runs:
            if run_data.hyperlink_url:
                self._add_hyperlink_run(paragraph, run_data)
                continue
            run = paragraph.add_run(run_data.text or "")
            self._apply_run_formatting(run, run_data)
            for media_item in run_data.embedded_media:
                self._add_media_to_paragraph(paragraph, media_item)

    def _apply_run_formatting(self, run, run_data) -> None:
        if run_data.bold is not None:
            run.bold = run_data.bold
        if run_data.italic is not None:
            run.italic = run_data.italic
        if run_data.underline is not None:
            run.underline = run_data.underline
        if getattr(run_data, "strikethrough", None):
            run.font.strike = True
        if run_data.font_name:
            run.font.name = run_data.font_name
        if run_data.font_size_pt is not None and run_data.font_size_pt > 0:
            run.font.size = Pt(run_data.font_size_pt)
        if run_data.color_rgb:
            try:
                hex_str = run_data.color_rgb.replace("#", "").strip()
                run.font.color.rgb = RGBColor(
                    int(hex_str[0:2], 16),
                    int(hex_str[2:4], 16),
                    int(hex_str[4:6], 16),
                )
            except Exception:
                pass
        if run_data.highlight_color:
            try:
                run.font.highlight_color = WD_COLOR_INDEX[run_data.highlight_color.upper(
                )]
            except (KeyError, AttributeError):
                pass

    def _add_hyperlink_run(self, paragraph, run_data) -> None:
        url = run_data.hyperlink_url or ""
        text = run_data.text or ""
        try:
            r_id = paragraph.part.relate_to(
                url, RT.HYPERLINK, is_external=True)
        except (AttributeError, KeyError, TypeError, ValueError):
            run = paragraph.add_run(text)
            self._apply_run_formatting(run, run_data)
            return

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        run_elem = self._build_hyperlink_run_element(run_data, text)
        self._append_text_to_oxml_run(run_elem, text)

        hyperlink.append(run_elem)
        # pylint: disable=protected-access
        paragraph._p.append(hyperlink)

    def _build_hyperlink_run_element(self, run_data, text: str):
        """Build oxml run element with hyperlink style overrides."""
        del text
        run_elem = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")

        hyperlink_blue = self._resolve_hyperlink_color(run_data)
        color_elem = OxmlElement("w:color")
        color_elem.set(qn(W_VAL_LITERAL), hyperlink_blue)
        rpr.append(color_elem)

        if run_data.underline is not False:
            u_elem = OxmlElement("w:u")
            u_elem.set(qn(W_VAL_LITERAL), "single")
            rpr.append(u_elem)

        if run_data.bold:
            rpr.append(OxmlElement("w:b"))
        if run_data.italic:
            rpr.append(OxmlElement("w:i"))
        if run_data.font_name:
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), run_data.font_name)
            r_fonts.set(qn("w:hAnsi"), run_data.font_name)
            rpr.append(r_fonts)
        if run_data.font_size_pt and run_data.font_size_pt > 0:
            half_pts = str(int(run_data.font_size_pt * 2))
            sz = OxmlElement("w:sz")
            sz.set(qn(W_VAL_LITERAL), half_pts)
            rpr.append(sz)
            sz_cs = OxmlElement("w:szCs")
            sz_cs.set(qn(W_VAL_LITERAL), half_pts)
            rpr.append(sz_cs)

        run_elem.append(rpr)
        return run_elem

    def _resolve_hyperlink_color(self, run_data) -> str:
        """Return hyperlink color hex (without #), defaulting to theme-friendly blue."""
        if not run_data.color_rgb:
            return "0563C1"
        return run_data.color_rgb.replace("#", "").strip()

    def _add_extracted_table(self, document: Document, table_data: ExtractedTable) -> None:
        if not table_data.rows:
            return

        column_count = max((len(row.cells)
                           for row in table_data.rows), default=0)
        if column_count == 0:
            return

        table = document.add_table(
            rows=len(table_data.rows), cols=column_count)
        self._apply_table_style(table, table_data)
        self._populate_docx_table(table, table_data)

    def _apply_table_style(self, table, table_data: ExtractedTable) -> None:
        """Apply extracted table style when available."""
        if not table_data.style:
            return
        try:
            table.style = table_data.style
        except KeyError:
            return

    def _populate_docx_table(self, table, table_data: ExtractedTable) -> None:
        """Populate a docx table recursively from extracted table data."""
        for row_index, row in enumerate(table_data.rows):
            for column_index, cell_data in enumerate(row.cells):
                target_cell = table.cell(row_index, column_index)
                merged_cell = self._merge_target_cell(
                    table,
                    target_cell,
                    cell_data,
                    row_index,
                    column_index,
                )
                self._populate_docx_cell(merged_cell, cell_data)

    def _merge_target_cell(self, table, target_cell, cell_data, row_index: int, column_index: int):
        """Merge table cells when colspan or rowspan is present."""
        colspan = getattr(cell_data, "colspan", None) or 1
        rowspan = getattr(cell_data, "rowspan", None) or 1
        if colspan == 1 and rowspan == 1:
            return target_cell

        end_row = min(row_index + rowspan - 1, len(table.rows) - 1)
        end_col = min(column_index + colspan - 1, len(table.columns) - 1)
        if end_row == row_index and end_col == column_index:
            return target_cell

        return target_cell.merge(table.cell(end_row, end_col))

    def _populate_docx_cell(self, cell, cell_data) -> None:
        """Populate a docx table cell with paragraphs and nested tables."""
        cell.text = ""
        paragraphs = list(getattr(cell_data, "paragraphs", []) or [])

        if paragraphs:
            self._populate_existing_cell_paragraph(
                cell.paragraphs[0], paragraphs[0])
            for paragraph_data in paragraphs[1:]:
                paragraph = cell.add_paragraph()
                self._populate_output_paragraph(paragraph, paragraph_data)
        elif getattr(cell_data, "text", None):
            cell.paragraphs[0].add_run(cell_data.text)

        for nested_table in getattr(cell_data, "tables", []) or []:
            self._add_nested_docx_table(cell, nested_table)

    def _populate_existing_cell_paragraph(self, paragraph, paragraph_data: ExtractedParagraph) -> None:
        """Populate the default paragraph already present in a table cell."""
        paragraph.text = ""
        self._populate_output_paragraph(paragraph, paragraph_data)

    def _add_nested_docx_table(self, cell, table_data: ExtractedTable) -> None:
        """Add a nested table inside a table cell."""
        if not table_data.rows:
            return

        column_count = max((len(row.cells)
                           for row in table_data.rows), default=0)
        if column_count == 0:
            return

        nested_table = cell.add_table(
            rows=len(table_data.rows), cols=column_count)
        self._apply_table_style(nested_table, table_data)
        self._populate_docx_table(nested_table, table_data)

    def _add_paragraph_block(self, document: Document, block: ParagraphBlock) -> None:
        if block.heading_level is not None:
            paragraph = document.add_heading(level=block.heading_level)
        else:
            paragraph = document.add_paragraph()

        run = paragraph.add_run(block.text)
        if block.bold is not None:
            run.bold = block.bold
        if block.italic is not None:
            run.italic = block.italic
        if block.underline is not None:
            run.underline = block.underline
        if block.font_name:
            run.font.name = block.font_name
        if block.font_size_pt is not None:
            run.font.size = Pt(block.font_size_pt)

        if block.alignment:
            paragraph.alignment = WD_ALIGN_PARAGRAPH[block.alignment]

    def _add_table_block(self, document: Document, block: TableBlock) -> None:
        if not block.rows:
            return

        column_count = max(len(row) for row in block.rows)
        if column_count == 0:
            return

        table = document.add_table(rows=len(block.rows), cols=column_count)
        if block.style:
            table.style = block.style

        for row_index, row in enumerate(block.rows):
            for column_index in range(column_count):
                text = row[column_index] if column_index < len(row) else ""
                table.cell(row_index, column_index).text = text

    def _resolve_paragraph_style_name(self, paragraph_data: ExtractedParagraph) -> str | None:
        if paragraph_data.is_numbered:
            return "List Number"

        if paragraph_data.is_bullet:
            return "List Bullet"

        if paragraph_data.numbering_format:
            fmt = paragraph_data.numbering_format.split(":", 1)[0].lower()
            if fmt == "bullet":
                return "List Bullet"
            return "List Number"

        if paragraph_data.style:
            return paragraph_data.style

        return None

    def _map_alignment(self, raw_alignment: str | None) -> WD_ALIGN_PARAGRAPH | None:
        if raw_alignment is None:
            return None

        normalized = raw_alignment.strip().upper()
        if normalized.startswith("LEFT"):
            return WD_ALIGN_PARAGRAPH.LEFT
        if normalized.startswith("CENTER"):
            return WD_ALIGN_PARAGRAPH.CENTER
        if normalized.startswith("RIGHT"):
            return WD_ALIGN_PARAGRAPH.RIGHT
        if normalized.startswith("JUSTIFY"):
            return WD_ALIGN_PARAGRAPH.JUSTIFY
        return None

    def _add_media_to_paragraph(self, paragraph, media_item: ExtractedMediaItem | str) -> None:
        """Insert an inline image using file path or in-payload base64 bytes."""
        local_file_path = (
            media_item.local_file_path
            if hasattr(media_item, "local_file_path")
            else media_item
        )
        base64_data = getattr(media_item, "base64_data", None) or getattr(
            media_item, "base64", None)

        width_emu = getattr(media_item, "width_emu", None)
        height_emu = getattr(media_item, "height_emu", None)

        try:
            run = paragraph.add_run()

            picture_source = None
            if base64_data:
                picture_source = BytesIO(base64.b64decode(base64_data))
            elif local_file_path:
                media_path = Path(local_file_path)
                if media_path.exists() and media_path.is_file():
                    picture_source = str(media_path)

            if picture_source is None:
                return

            if width_emu and height_emu:
                run.add_picture(picture_source, width=Emu(
                    width_emu), height=Emu(height_emu))
            elif width_emu:
                run.add_picture(picture_source, width=Emu(width_emu))
            else:
                run.add_picture(picture_source, width=Inches(2.5))
        except Exception:
            # Ignore invalid/unsupported image data and continue with text content.
            return
